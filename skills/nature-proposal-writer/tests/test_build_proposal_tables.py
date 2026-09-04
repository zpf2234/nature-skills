from __future__ import annotations

import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SCRIPT = ROOT / "scripts" / "build_proposal_docx.py"


class ProposalTableTests(unittest.TestCase):
    def test_alignment_separator_is_not_exported_as_a_data_row(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "proposal.md"
            output = root / "proposal.docx"
            source.write_text(
                "# Proposal\n\n"
                "| Item | Score |\n"
                "| :--- | ---: |\n"
                "| Baseline | 0.91 |\n",
                encoding="utf-8",
            )

            result = subprocess.run(
                [sys.executable, str(SCRIPT), str(source), str(output)],
                text=True,
                capture_output=True,
                check=False,
            )

            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            document = Document(output)
            self.assertEqual(len(document.tables), 1)
            self.assertEqual(len(document.tables[0].rows), 2)
            self.assertEqual(
                [[cell.text for cell in row.cells] for row in document.tables[0].rows],
                [["Item", "Score"], ["Baseline", "0.91"]],
            )


if __name__ == "__main__":
    unittest.main()
