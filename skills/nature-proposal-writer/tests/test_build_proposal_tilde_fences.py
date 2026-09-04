from __future__ import annotations

import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SCRIPT = ROOT / "scripts" / "build_proposal_docx.py"


class ProposalTildeFenceTests(unittest.TestCase):
    def test_tilde_fence_is_rendered_as_code(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "proposal.md"
            output = root / "proposal.docx"
            source.write_text(
                "# Proposal\n\n~~~python\nresult = 42\n~~~\n",
                encoding="utf-8",
            )

            result = subprocess.run(
                [sys.executable, str(SCRIPT), str(source), str(output)],
                text=True,
                capture_output=True,
                check=False,
            )

            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            document = Document(output)
            code_runs = [
                run
                for paragraph in document.paragraphs
                for run in paragraph.runs
                if run.text == "result = 42"
            ]
            self.assertEqual(len(code_runs), 1)
            self.assertEqual(code_runs[0].font.name, "Consolas")
            self.assertNotIn("~~~", "\n".join(p.text for p in document.paragraphs))


if __name__ == "__main__":
    unittest.main()
